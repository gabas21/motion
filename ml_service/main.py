from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, field_validator
import numpy as np
import requests
import io
import pypdf
import os
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from intent_classifier import IntentClassifier, parse_indonesian_date, extract_task_name

app = FastAPI(
    title="Motion ML & RAG Engine",
    description="Microservice untuk Rule-Based Analytics dan PDF Embeddings Generator",
    version="1.1"
)

classifier = IntentClassifier()
classifier.train()

# ─────────────────────────────────────────────────────────────────────────────
# Schema Data
# ─────────────────────────────────────────────────────────────────────────────

class TaskItem(BaseModel):
    id: str
    status: str
    time_estimate_minutes: int
    due_date: str = None
    completed_at: str = None
    scheduled_start: str = None

class MoodleAssignmentItem(BaseModel):
    id: str
    courseId: str
    courseName: str
    name: str
    dueDate: str = None
    submissionStatus: str = None # new, draft, submitted

class MLInputData(BaseModel):
    tasks: list[TaskItem]
    moodleAssignments: list[MoodleAssignmentItem] = []

    @field_validator('tasks')
    @classmethod
    def validate_tasks_limit(cls, v):
        if len(v) > 500:
            raise ValueError("Maksimum 500 task per request untuk menjaga performa server")
        return v

class RouteRequest(BaseModel):
    text: str

# ─────────────────────────────────────────────────────────────────────────────
# Health Check
# ─────────────────────────────────────────────────────────────────────────────

@app.get("/health")
def health_check():
    return {
        "status": "healthy",
        "version": "1.1",
        "service": "motion-ml"
    }

@app.post("/agent/route")
def route_agent_query(data: RouteRequest):
    if not data.text or not data.text.strip():
        return {
            "intent": "general_chat",
            "confidence": 1.0,
            "entities": {}
        }
    
    intent, confidence = classifier.predict(data.text)
    
    entities = {}
    if intent == "schedule_event":
        entities["date"] = parse_indonesian_date(data.text)
        entities["task_name"] = extract_task_name(data.text)
    elif intent in ["sync_welearn", "check_deadline"]:
        parsed_date = parse_indonesian_date(data.text)
        if parsed_date:
            entities["date"] = parsed_date
            
    return {
        "intent": intent,
        "confidence": confidence,
        "entities": entities
    }

# ─────────────────────────────────────────────────────────────────────────────
# 1. ANALYTICS: Burnout Risk Calculator (Rule-Based — jelas dan jujur)
#
# Sebelumnya menggunakan LogisticRegression dari sklearn, tapi model di-train
# hanya dengan 2 data point dan langsung di-overwrite koefisiennya secara manual.
# Itu bukan machine learning — itu matematika manual dengan wrapper sklearn.
# Versi ini menggunakan pendekatan rule-based yang lebih jujur, lebih cepat,
# dan hasil yang identik secara matematis.
# ─────────────────────────────────────────────────────────────────────────────

@app.post("/predict/burnout")
def predict_burnout(data: MLInputData):
    if not data.tasks:
        return {
            "score": 12.5,
            "status": "Low",
            "description": "Beban kerja aman. Pertahankan keseimbangan hidup!"
        }

    import pandas as pd

    df = pd.DataFrame([t.model_dump() for t in data.tasks])

    total_tasks = len(df)
    pending_tasks = df[df['status'] != 'completed']
    completed_tasks = df[df['status'] == 'completed']

    # Hitung fitur
    overdue_count = 0
    now = pd.Timestamp.now(tz='UTC')

    for _, row in pending_tasks.iterrows():
        if row['due_date']:
            due_dt = pd.to_datetime(row['due_date'])
            if due_dt.tzinfo is None:
                due_dt = due_dt.tz_localize('UTC')
            if now > due_dt:
                overdue_count += 1

    overdue_ratio = overdue_count / total_tasks if total_tasks > 0 else 0.0
    total_estimate_min = pending_tasks['time_estimate_minutes'].sum() if len(pending_tasks) > 0 else 0

    # 480 menit = 8 jam kerja aktif ideal per hari
    workload_density = float(total_estimate_min) / 480.0

    # Hitung pengerjaan larut malam (12 AM - 5 AM)
    midnight_study_count = 0
    if len(completed_tasks) > 0:
        for _, row in completed_tasks.iterrows():
            if row['completed_at']:
                comp_hour = pd.to_datetime(row['completed_at']).hour
                if 0 <= comp_hour < 5:
                    midnight_study_count += 1

    # Batasi fitur begadang maksimum 4
    mid_study_feature = min(float(midnight_study_count), 4.0)

    # ── Formula Rule-Based (bobot identik dengan koefisien model sebelumnya) ──
    # Bobot: overdue_ratio=4.5, mid_study=0.8, workload_density=1.2, intercept=-2.2
    # Sigmoid diterapkan secara manual untuk menghasilkan probabilitas 0-1.
    raw_score = (overdue_ratio * 4.5) + (mid_study_feature * 0.8) + (workload_density * 1.2) - 2.2
    import math
    probability = 1.0 / (1.0 + math.exp(-raw_score))
    score_percent = probability * 100.0

    # Batasi persentase logis
    score_percent = max(min(score_percent, 98.0), 5.0)

    status = "Low"
    desc = "Beban kerja Anda terpantau aman dan seimbang. Sangat baik untuk kesehatan mental Anda!"

    if score_percent >= 70.0:
        status = "High"
        desc = "RISIKO BURNOUT TINGGI! Terlalu banyak beban tugas tertunda dan pola begadang. Segera istirahat dan hubungi Hermes AI untuk menjadwal ulang!"
    elif score_percent >= 35.0:
        status = "Moderate"
        desc = "Risiko stres sedang terdeteksi. Cobalah mencicil tugas overdue dan pastikan tidur cukup malam ini."

    return {
        "score": float(score_percent),
        "status": status,
        "description": desc
    }

# ─────────────────────────────────────────────────────────────────────────────
# 2. ANALYTICS: Golden Hours Finder
#
# Sebelumnya menggunakan KMeans dengan n_clusters=1, yang ekuivalen dengan
# menghitung centroid = rata-rata. Versi ini menggunakan np.mean() secara
# langsung yang lebih efisien dan transparan.
# ─────────────────────────────────────────────────────────────────────────────

@app.post("/predict/golden-hours")
def predict_golden_hours(data: MLInputData):
    if not data.tasks:
        return {"peakDay": "Butuh Data", "peakHourRange": "Selesaikan tugas dahulu", "confidence": "0%"}

    import pandas as pd

    df = pd.DataFrame([t.model_dump() for t in data.tasks])
    completed = df[df['status'] == 'completed'].copy()

    if len(completed) < 3:
        return {
            "peakDay": "Butuh Data",
            "peakHourRange": f"Selesaikan minimal {3 - len(completed)} tugas lagi",
            "confidence": "10%"
        }

    # Extract Hour dan DayOfWeek
    completed['completed_at'] = pd.to_datetime(completed['completed_at'])
    completed['hour'] = completed['completed_at'].dt.hour
    completed['dayofweek'] = completed['completed_at'].dt.dayofweek

    X = completed[['hour', 'dayofweek']].values

    # Hitung pusat distribusi menggunakan rata-rata (setara KMeans n_clusters=1)
    center_hour = float(np.mean(X[:, 0]))
    center_day = float(np.mean(X[:, 1]))

    days_indo = ["Minggu", "Senin", "Selasa", "Rabu", "Kamis", "Jumat", "Sabtu"]
    best_day = days_indo[int(round(center_day)) % 7]
    best_hour = int(round(center_hour))

    start_hour = (best_hour - 1) % 24
    end_hour = (best_hour + 1) % 24

    # Hitung keyakinan model (Confidence) berdasarkan sebaran data
    distances = np.sqrt(np.sum((X - np.array([center_hour, center_day])) ** 2, axis=1))
    mean_distance = float(np.mean(distances)) if len(distances) > 0 else 1.0

    # Skala confidence terbalik terhadap jarak rata-rata ke pusat
    conf = 100.0 - (mean_distance * 10.0)
    conf = max(min(conf, 95.0), 15.0)

    return {
        "peakDay": best_day,
        "peakHourRange": f"{start_hour:02d}:00 - {end_hour:02d}:00 WIB",
        "confidence": f"{conf:.0f}%"
    }

# ─────────────────────────────────────────────────────────────────────────────
# 2b. ANALYTICS: Graduation Risk Index (GRI) Predictor
# ─────────────────────────────────────────────────────────────────────────────

@app.post("/predict/graduation-risk")
def predict_graduation_risk(data: MLInputData):
    import pandas as pd
    now = pd.Timestamp.now(tz='UTC')
    total_assigns = len(data.moodleAssignments)

    if total_assigns == 0:
        # Jika tidak ada tugas WeLearn, hitung berbasis tugas reguler
        total_tasks = len(data.tasks)
        if total_tasks == 0:
            return {
                "score": 10.0,
                "status": "Low",
                "description": "Belum ada tugas terdaftar. Risiko akademik terpantau sangat rendah."
            }
        
        overdue_tasks = 0
        for t in data.tasks:
            if t.status != 'completed' and t.due_date:
                due_dt = pd.to_datetime(t.due_date)
                if due_dt.tzinfo is None:
                    due_dt = due_dt.tz_localize('UTC')
                if now > due_dt:
                    overdue_tasks += 1
        
        overdue_ratio = overdue_tasks / total_tasks
        academic_score = 1.0 - overdue_ratio
        consistency_score = 1.0
    else:
        overdue_assigns = 0
        missed_quizzes = 0
        quizzes_count = 0
        completed_assigns = 0

        for a in data.moodleAssignments:
            is_quiz = any(q in a.name.lower() for q in ["quiz", "kuis", "uts", "uas", "praktikum"])
            if is_quiz:
                quizzes_count += 1

            if a.submissionStatus != 'submitted':
                if a.dueDate:
                    due_dt = pd.to_datetime(a.dueDate)
                    if due_dt.tzinfo is None:
                        due_dt = due_dt.tz_localize('UTC')
                    if now > due_dt:
                        overdue_assigns += 1
                        if is_quiz:
                            missed_quizzes += 1
            else:
                completed_assigns += 1

        overdue_ratio = overdue_assigns / total_assigns
        quiz_miss_ratio = missed_quizzes / quizzes_count if quizzes_count > 0 else overdue_ratio
        
        # Academic Score: prioritas kelengkapan tugas (60% tugas reguler, 40% kuis/praktikum)
        academic_score = 1.0 - (overdue_ratio * 0.6 + quiz_miss_ratio * 0.4)

        # Procrastination checking on completed regular tasks
        procrastination_count = 0
        completed_tasks = [t for t in data.tasks if t.status == 'completed']
        for t in completed_tasks:
            if t.completed_at and t.due_date:
                comp_dt = pd.to_datetime(t.completed_at)
                due_dt = pd.to_datetime(t.due_date)
                if comp_dt.tzinfo is None:
                    comp_dt = comp_dt.tz_localize('UTC')
                if due_dt.tzinfo is None:
                    due_dt = due_dt.tz_localize('UTC')
                
                # Cek jika selesai dalam 3 jam sebelum deadline
                diff_hours = (due_dt - comp_dt).total_seconds() / 3600.0
                if 0 <= diff_hours <= 3.0:
                    procrastination_count += 1

        procrastination_ratio = procrastination_count / len(completed_tasks) if len(completed_tasks) > 0 else 0.0
        consistency_score = 1.0 - procrastination_ratio

    # Hitung Graduation Risk Index (GRI)
    # Bobot: Academic Score 70%, Consistency Score 30%
    gri_score = (1.0 - (max(0.0, academic_score) * 0.7 + max(0.0, consistency_score) * 0.3)) * 100.0
    gri_score = max(min(gri_score, 98.0), 5.0)

    status = "Low"
    desc = "Risiko akademik rendah. Pola pengerjaan tugas dan tingkat kelulusan tepat waktu Anda berada di jalur aman!"

    if gri_score >= 60.0:
        status = "High"
        desc = "RISIKO AKADEMIK TINGGI! Terlalu banyak tugas WeLearn terlambat/terbengkalai. Segera cicil tugas kuliah agar kelulusan tepat waktu aman."
    elif gri_score >= 30.0:
        status = "Moderate"
        desc = "Risiko akademik sedang terdeteksi. Beberapa tugas WeLearn terlewat. Gunakan AI Scheduler untuk menyusun ulang jadwal belajar Anda."

    return {
        "score": float(gri_score),
        "status": status,
        "description": desc
    }

# ─────────────────────────────────────────────────────────────────────────────
# 3. RAG: Chunking & Embeddings Generator (Gemini model API)
# ─────────────────────────────────────────────────────────────────────────────

@app.post("/documents/embed")
async def generate_document_embeddings(
    file: UploadFile = File(...),
    gemini_api_key: str = Form(...)
):
    # 1. Baca isi berkas
    contents = await file.read()
    filename = file.filename
    raw_text = ""

    if filename.endswith(".pdf"):
        # Ekstrak teks dari PDF menggunakan PyPDF
        try:
            pdf_file = io.BytesIO(contents)
            reader = pypdf.PdfReader(pdf_file)
            text_list = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_list.append(text)
            raw_text = "\n".join(text_list)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Gagal membaca PDF: {str(e)}")
    else:
        # Asumsikan text/markdown biasa
        try:
            raw_text = contents.decode("utf-8")
        except Exception:
            try:
                raw_text = contents.decode("latin-1")
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Gagal membaca teks berkas: {str(e)}")

    if not raw_text.strip():
        raise HTTPException(status_code=400, detail="File kosong atau tidak mengandung teks terdeteksi.")

    # 2. Text Chunking (Pemotongan)
    # Ukuran 500 karakter, overlap 100 karakter
    chunk_size = 500
    overlap = 100
    chunks = []

    start = 0
    while start < len(raw_text):
        end = start + chunk_size
        chunk = raw_text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start += chunk_size - overlap

    if not chunks:
        raise HTTPException(status_code=400, detail="Gagal membagi dokumen menjadi potongan-potongan teks.")

    # Batasi jumlah chunk untuk mencegah abuse
    max_chunks = 200
    if len(chunks) > max_chunks:
        chunks = chunks[:max_chunks]
        print(f"[ML-Embed-Warn] Dokumen dipotong menjadi {max_chunks} chunk pertama (total asli: {len(chunks)})")

    # 3. Kirim pecahan teks ke Gemini Embeddings API (text-embedding-004) secara berurutan
    embed_results = []
    headers = {"Content-Type": "application/json"}

    url = f"https://generativelanguage.googleapis.com/v1beta/models/text-embedding-004:embedContent?key={gemini_api_key}"

    for i, chunk in enumerate(chunks):
        payload = {
            "content": {
                "parts": [{
                    "text": chunk
                }]
            }
        }

        try:
            response = requests.post(url, headers=headers, json=payload, timeout=10)
            if response.status_code == 200:
                res_data = response.json()
                vector = res_data.get("embedding", {}).get("values", [])
                if vector:
                    embed_results.append({
                        "content": chunk,
                        "embedding": vector # List of float (768 dimensi)
                    })
            else:
                print(f"[ML-Embed-Error] Gagal generate chunk {i}: {response.status_code} - {response.text}")
        except Exception as e:
            print(f"[ML-Embed-Error] Exception di chunk {i}: {str(e)}")

    if not embed_results:
        raise HTTPException(
            status_code=500,
            detail="Gagal menghasilkan vektor embedding dari Gemini API. Pastikan kunci API Gemini valid."
        )

    return {
        "documentName": filename,
        "chunksCount": len(embed_results),
        "data": embed_results
    }

# ─────────────────────────────────────────────────────────────────────────────
# 4. DOCX: Document Generator (python-docx)
# ─────────────────────────────────────────────────────────────────────────────

class DocxRequest(BaseModel):
    title: str
    content: str

@app.post("/documents/generate-docx")
def generate_docx(data: DocxRequest):
    doc = docx.Document()

    # Set standard margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title styling
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(data.title)
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(16)
    title_run.bold = True
    title_p.paragraph_format.space_after = Pt(24)

    lines = data.content.split("\n")

    # Helper to parse inline formatting (bold/italic) in a paragraph
    def add_formatted_text(paragraph, text, is_list=False):
        import re
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                run = paragraph.add_run(part)
            run.font.name = 'Arial'
            run.font.size = Pt(11)

    # Process lines
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Detect markdown tables
        if line.startswith("|") and i + 1 < len(lines) and (lines[i+1].strip().startswith("|---") or lines[i+1].strip().startswith("|-")):
            table_lines = []
            table_lines.append(line)
            # Skip separator line
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1

            if len(table_lines) > 0:
                rows_data = []
                for tl in table_lines:
                    cols = [c.strip() for c in tl.split("|")[1:-1]]
                    rows_data.append(cols)

                if rows_data:
                    num_cols = max(len(row) for row in rows_data) if rows_data else 1
                    num_rows = len(rows_data)
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Light Shading Accent 1'

                    for r_idx, row in enumerate(rows_data):
                        # Pad shorter rows and slice longer rows to match num_cols
                        if len(row) < num_cols:
                            row = row + [""] * (num_cols - len(row))
                        elif len(row) > num_cols:
                            row = row[:num_cols]

                        for c_idx, val in enumerate(row):
                            cell = table.cell(r_idx, c_idx)
                            cell.text = ""
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(4)
                            p.paragraph_format.space_after = Pt(4)
                            run = p.add_run(val)
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)
                            if r_idx == 0:
                                run.bold = True
                    p_after = doc.add_paragraph()
                    p_after.paragraph_format.space_before = Pt(12)
            continue

        # Headings
        if line.startswith("# "):
            h = doc.add_paragraph()
            run = h.add_run(line[2:])
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.bold = True
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(6)
        elif line.startswith("## "):
            h = doc.add_paragraph()
            run = h.add_run(line[3:])
            run.font.name = 'Arial'
            run.font.size = Pt(13)
            run.bold = True
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
        elif line.startswith("### "):
            h = doc.add_paragraph()
            run = h.add_run(line[4:])
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.bold = True
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(4)
        elif line.startswith("#### "):
            h = doc.add_paragraph()
            run = h.add_run(line[5:])
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.bold = True
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
        # Bullet lists
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            add_formatted_text(p, line[2:], is_list=True)
        # Numbered lists
        elif line[0].isdigit() and (". " in line or ") " in line):
            dot_idx = line.find(". ")
            paren_idx = line.find(") ")
            split_idx = dot_idx if (dot_idx != -1 and (paren_idx == -1 or dot_idx < paren_idx)) else paren_idx
            if split_idx != -1:
                p = doc.add_paragraph(style='List Number')
                p.paragraph_format.space_after = Pt(4)
                add_formatted_text(p, line[split_idx+2:], is_list=True)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = 1.15
                add_formatted_text(p, line)
        # Regular paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15
            add_formatted_text(p, line)

        i += 1

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=jawaban.docx"}
    )
